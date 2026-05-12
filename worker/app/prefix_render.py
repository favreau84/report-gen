"""Rendu de la convention par préfixe : substitution `<prefix>xxx` → valeur JSON.

Le préfixe est paramétrable (par ex. `li_`, `tag_`, `bal_`). La résolution est
case-insensitive.

Limites de la version MVP :
- Substitution texte uniquement.
- Si la clé n'est pas dans le JSON, la balise est laissée telle quelle (utile
  pour repérer les manques dans le PDF final).
- Les **blocs** `<prefix>X_DEBUT … FIN` et `_start … _stop` ne sont pas encore
  évalués comme des conditions — les marqueurs apparaîtront tels quels.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph


def render_with_prefix(
    src_docx: Path,
    dst_docx: Path,
    payload: dict[str, Any],
    prefix: str,
) -> None:
    """Charge `src_docx`, substitue tous les `<prefix>xxx` à partir de `payload`,
    et écrit le résultat dans `dst_docx`.
    """
    prefix = (prefix or "").strip().lower()
    if not prefix:
        Document(str(src_docx)).save(str(dst_docx))
        return

    doc = Document(str(src_docx))
    flat = _flatten_keys(payload)
    skip_len = len(prefix)

    pattern = re.compile(rf"(?i){re.escape(prefix)}(?:(?!{re.escape(prefix)})\w)+")

    def replace(m: re.Match[str]) -> str:
        tok = m.group(0)
        key = tok[skip_len:].lower()
        return flat.get(key, tok)

    for para in _iter_paragraphs(doc):
        _substitute_paragraph(para, pattern, replace)

    doc.save(str(dst_docx))


def _iter_paragraphs(doc: Document) -> "list[Paragraph]":
    paragraphs: list[Paragraph] = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _substitute_paragraph(
    para: Paragraph,
    pattern: re.Pattern[str],
    replace,
) -> None:
    """Substitue les balises dans un paragraphe en collapsant les runs.

    Le formatage par-run est perdu sur les paragraphes substitués, mais le
    formatage paragraph-level (style, alignement) est conservé. Acceptable
    pour le MVP.
    """
    text = para.text
    if not pattern.search(text):
        return
    new_text = pattern.sub(replace, text)
    if new_text == text:
        return
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def _flatten_keys(payload: dict[str, Any], prefix: str = "") -> dict[str, str]:
    """Aplatit un JSON imbriqué en dot-notation, en lowercase."""
    out: dict[str, str] = {}
    for k, v in payload.items():
        key = f"{prefix}{str(k).lower()}"
        if isinstance(v, dict):
            out.update(_flatten_keys(v, prefix=key + "."))
        elif isinstance(v, list):
            out[key] = ", ".join(str(x) for x in v)
        elif v is None:
            out[key] = ""
        else:
            out[key] = str(v)
    return out
